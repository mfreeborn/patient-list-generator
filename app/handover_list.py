import logging
from datetime import datetime

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn
from docx.shared import Cm, Pt
from docx.table import Table

from app import enums, models, teams, trakcare

logger = logging.getLogger("PLG")


class HandoverTable:
    def __init__(self, table: Table):
        self._table = table
        self._new_patient_indices = set()

    def __getattr__(self, attr):
        return getattr(self._table, attr)

    def clear(self, keep_headers: bool = True) -> None:
        logger.debug("Removing old rows from original patient list table")
        for row in self.rows[keep_headers:]:
            tr = row._tr
            self._tbl.remove(tr)

    def add_ward_header_row(self, ward: enums.Ward) -> None:
        ward_header_row = self.add_row()
        ward_header_row.cells[0].merge(ward_header_row.cells[-1])
        ward_header_row.cells[0].text = ward.value

    def add_patient_row(self, patient: models.Patient) -> None:
        new_patient_row = self.add_row()
        new_patient_row.patient = patient

        new_patient_row.cells[0].text = patient.bed
        new_patient_row.cells[1].text = patient.patient_details
        new_patient_row.cells[2].text = patient.reason_for_admission or ""
        new_patient_row.cells[3].text = patient.jobs
        new_patient_row.cells[4].text = patient.edd
        new_patient_row.cells[5].text = patient.ds
        new_patient_row.cells[6].text = patient.tta
        new_patient_row.cells[7].text = patient.bloods

        if patient.is_new:
            self._new_patient_indices.add(new_patient_row._index)

    def format(self) -> None:
        # center table within the page
        self.alignment = WD_TABLE_ALIGNMENT.RIGHT

        # set column width for the patient details
        # 3cm is a good default width to fit dob and age on one line
        for cell in self.columns[1].cells:
            cell.width = Cm(3)

        # format the rows and cells
        for row in self.rows:
            # format the column headers and set header row to repeat
            if row.cells[0].text.lower() == "bed":
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.underline = True
                tr = row._tr
                trPr = tr.get_or_add_trPr()
                tblHeader = OxmlElement("w:tblHeader")
                tblHeader.set(qn("w:val"), "true")
                trPr.append(tblHeader)

            # format the ward headers
            if row.cells[0].text == row.cells[1].text:
                shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="EEEEEE"/>')
                for cell in row.cells:
                    cell._tc.get_or_add_tcPr().append(shading_elm)
                    # keep this header row with on the same page as the next row
                    for paragraph in cell.paragraphs:
                        paragraph.paragraph_format.keep_with_next = True

            # format new patients as bold
            if row._index in self._new_patient_indices:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        for run in para.runs:
                            run.bold = True

            # format all cells in the table
            for cell in row.cells:
                cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                for paragraph in cell.paragraphs:
                    formatter = paragraph.paragraph_format
                    formatter.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                    formatter.space_before = Pt(2)
                    formatter.space_after = Pt(2)

            # prevent table rows from splitting across page breaks
            no_split = parse_xml(f"<w:cantSplit {nsdecls('w')}/>")
            row._tr.get_or_add_trPr().append(no_split)


class HandoverList:
    """The primary object representing the Word document continaing the team's list of patients."""

    def __init__(self, team: teams.Team, file_path):
        logger.debug("Instantiating HandoverList using %s as a base list", file_path)
        self.doc = Document(docx=file_path)
        self.team = team
        self.patients = self._parse_patients()

        # set the default document format
        style = self.styles["Normal"]
        font = style.font
        font.name = "Calibri"
        font.size = Pt(8)
        sections = self.doc.sections

        # set narrow margins
        for section in sections:
            section.top_margin = Cm(1.25)
            section.bottom_margin = Cm(1.25)
            section.left_margin = Cm(1.25)
            section.right_margin = Cm(1.25)

    def __getattr__(self, attr):
        """Pass any unresolved attribute accesses down to the underlying docx.Document instance."""
        return getattr(self.doc, attr)

    def _parse_patients(self) -> models.PatientList:
        """Parse the patients from the Word document table into a PatientList."""
        patient_list = models.PatientList(home_ward=self.team.home_ward)

        for row in self._handover_table.rows:
            # skip past the column headers
            if row.cells[0].text.lower() == "bed":
                continue

            # skip past the ward name headers or empty rows
            if row.cells[0].text == row.cells[1].text:
                continue
            try:
                pt = models.Patient.from_table_row(row)
            except ValueError as e:
                logger.exception(e)
                raise

            patient_list.append(pt)

        logger.debug("%d patients found on the input patient list", len(patient_list))
        return patient_list

    @property
    def patient_count(self) -> int:
        return len(self.patients)

    @property
    def new_patient_count(self) -> int:
        return sum(patient.is_new for patient in self.patients)

    @property
    def _handover_table(self) -> HandoverTable:
        """Return the underlying Word document table."""
        return HandoverTable(self.tables[0])

    def _update_patients(self) -> None:
        """Bring the internal PatientList object up to date with TrakCare."""
        updated_list = models.PatientList(home_ward=self.team.home_ward)
        current_trakcare_patients = trakcare.get_trakcare_patients(self.team)

        for trakcare_patient in current_trakcare_patients:
            if trakcare_patient not in self.patients:
                # patient must be new to the team
                trakcare_patient.is_new = True
                updated_list.append(trakcare_patient)
            else:
                # patient must be on the original list, therefore merge the
                # jobs, EDD etc into the trakcare patient. By merging into
                # the trakcare patient, we also ensure that their location is fully
                # up to date incase they were moved since the previous list.
                trakcare_patient.merge(self.patients[trakcare_patient.nhs_number])
                updated_list.append(trakcare_patient)

        updated_list.sort()
        self.patients = updated_list
        logger.debug("Patient list updated: %s", self.patients)

    def _update_handover_table(self) -> None:
        """Create a fresh Word table with the current PatientList."""
        logger.debug("Updating Microsoft Word patient list")
        table = self._handover_table

        # clear the existing table back to the column headers
        table.clear()

        logger.debug("Adding new and updated patient rows to patient list table")
        current_ward = None
        for patient in self.patients:
            if patient.location.ward != current_ward:
                # we're at the first patient on a different ward, therefore add
                # a subheader row with the name of the new ward
                current_ward = patient.location.ward
                table.add_ward_header_row(current_ward)

            table.add_patient_row(patient)

        # apply formatting to the newly updated table
        logger.debug("Applying document formatting")
        table.format()

    def _update_list_metadata(self) -> None:
        """Update additonal metadata, such as the footer."""
        # put the patient count in the footer
        footer = self.sections[0].footer
        footer.footer_distance = Cm(1.25)
        footer.paragraphs[0].text = (
            f"{self.patient_count} patients ({self.new_patient_count} new)\t"
            f"\t"
            f"Generated at {datetime.now():%H:%M %d/%m/%Y}"
        )

    def update(self) -> None:
        """Update the HandoverList patient table.

        Comprises of 3 phases:
            1) Update the in-memory 'self.patients' PatientList using CareFlow +/- TrakCare
            2) Update the underlying table in the Word document with the updated 'self.patients'
            3) Update any addtional metadata pertaining to the HandoverList e.g. footer information
        """
        self._update_patients()
        self._update_handover_table()
        self._update_list_metadata()
