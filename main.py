import datetime

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls
from docx.shared import Pt
from docx.table import _Row

from careflow import get_careflow_patients
from enums import TeamName
from pt_list import get_list_patients
from teams import TEAMS, Team
from utils import Patient, PatientList, remove_table_row


def remove_olds(orig_list: PatientList, careflow_list: PatientList) -> PatientList:
    # iterate over the original list and keep only those who also appear in the careflow list

    new_list = PatientList(home_ward=orig_list.home_ward)
    for pt in orig_list:
        if pt in careflow_list:
            new_list.append(pt)

    return new_list


def add_news(orig_list: PatientList, careflow_list: PatientList) -> PatientList:
    # iterate over the careflow list and, if the patient isn't in the original list, add them
    for pt in careflow_list:
        if pt not in orig_list:
            pt.is_new = True
            orig_list.append(pt)

    return orig_list


def add_patient(patient: Patient, row: _Row):
    row.cells[0].text = patient.bed
    row.cells[1].text = patient.patient_details
    row.cells[2].text = patient.reason_for_admission
    row.cells[3].text = patient.jobs
    row.cells[4].text = patient.edd
    row.cells[5].text = patient.ds
    row.cells[6].text = patient.tta
    row.cells[7].text = patient.bloods


def update_msword_list(doc: Document, patient_list: PatientList, team_name: str):
    table = doc.tables[0]

    for row in table.rows[1:]:
        # skip over the column headers row alone
        remove_table_row(table, row)

    # set the default document format
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Arial"
    font.size = Pt(8)

    # generate the new patient list table
    current_ward = None
    for pt in patient_list:
        if pt.location.ward != current_ward:
            # we're at the first patient on a different ward, therefore add
            # a subheader row with the name of the new ward
            ward_header_row = table.add_row()
            ward_header_row.cells[0].merge(ward_header_row.cells[-1])
            ward_header_row.cells[0].text = pt.location.ward.value

            # format the ward header row
            shading_elm = parse_xml(r'<w:shd {} w:fill="EEEEEE"/>'.format(nsdecls("w")))
            for cell in ward_header_row.cells:
                for paragraph in cell.paragraphs:
                    formatter = paragraph.paragraph_format
                    formatter.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                    formatter.space_after = Pt(0)

                cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                cell._tc.get_or_add_tcPr().append(shading_elm)

            current_ward = pt.location.ward

        row = table.add_row()
        add_patient(pt, row)

        # format the patient row
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                formatter = paragraph.paragraph_format
                formatter.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                formatter.space_after = Pt(0)
                if pt.is_new:
                    for run in paragraph.runs:
                        run.bold = True

            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

    doc.save(f"{datetime.date.today():%d-%m-%y}_{team_name.lower()}.docm")


def generate_patient_list(team: Team, doc: Document):
    # first create two PatientList objects - one containing all the patients on CareFlow,
    # and the second containing all the patients on the previous day's paper list
    careflow_pts = get_careflow_patients(team=team)
    list_pts = get_list_patients(team=team)

    # start by taking the paper list and removing all of the patients which appear
    # here but not on the careflow list (i.e. patient has been discharged or moved team)
    new_list = remove_olds(list_pts, careflow_pts)

    # add the new patients to the list, which are those who appear on the careflow list,
    # but not the paper list
    new_list = add_news(new_list, careflow_pts)

    # sort the patients prior to inserting them into a Word table
    new_list.sort()

    # update and save the new list under today's date (will silently overwrite any
    # pre-existing file with the same name)
    update_msword_list(doc, new_list, team.name.value)


if __name__ == "__main__":
    team = TEAMS[TeamName.ELAMIN]
    doc = Document("31.12.2019.docm")

    generate_patient_list(team, doc)
