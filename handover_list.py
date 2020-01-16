from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls
from docx.shared import Pt, Inches

from careflow import get_careflow_patients
from enums import Ward
from patient import Patient, PatientList
from teams import Team


class HandoverTable:
    def __init__(self, table):
        self._table = table
        self._new_patient_indices = set()

    def __getattr__(self, attr):
        return getattr(self._table, attr)

    def clear(self, keep_headers=True):
        for row in self.rows[keep_headers:]:
            tr = row._tr
            self._tbl.remove(tr)

    def add_ward_header_row(self, ward: Ward):
        ward_header_row = self.add_row()
        ward_header_row.cells[0].merge(ward_header_row.cells[-1])
        ward_header_row.cells[0].text = ward.value

    def add_patient_row(self, patient: Patient):
        new_patient_row = self.add_row()
        new_patient_row.patient = patient

        new_patient_row.cells[0].text = patient.bed
        new_patient_row.cells[1].text = patient.patient_details
        new_patient_row.cells[2].text = patient.reason_for_admission
        new_patient_row.cells[3].text = patient.jobs
        new_patient_row.cells[4].text = patient.edd
        new_patient_row.cells[5].text = patient.ds
        new_patient_row.cells[6].text = patient.tta
        new_patient_row.cells[7].text = patient.bloods

        if patient.is_new:
            self._new_patient_indices.add(new_patient_row._index)

    def format(self):
        for row in self.rows:
            # format the column headers
            if row.cells[0].text.lower() == "bed":
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.underline = True

            # format the ward headers
            if row.cells[0].text == row.cells[1].text:
                shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="EEEEEE"/>')
                for cell in row.cells:
                    cell._tc.get_or_add_tcPr().append(shading_elm)
                    # keep this header row with on the same page as the next row
                    for paragraph in cell.paragraphs:
                        paragraph.paragraph_format.keep_with_next = True

            # format new patients as bold
            if row._index in self._new_patient_indices:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        for run in para.runs:
                            run.bold = True

            # format all cells in the table
            for cell in row.cells:
                cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                for paragraph in cell.paragraphs:
                    formatter = paragraph.paragraph_format
                    formatter.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                    formatter.space_before = Pt(2)
                    formatter.space_after = Pt(2)

            # prevent rows from splitting across page breaks
            no_split = parse_xml(f"<w:cantSplit {nsdecls('w')}/>")
            row._tr.get_or_add_trPr().append(no_split)


class HandoverList:
    def __init__(self, team: Team, file_path=None):
        self.doc = Document(docx=file_path)
        self.team = team
        self.patients: PatientList = self._parse_patients()

        # set the default document format
        style = self.styles["Normal"]
        font = style.font
        font.name = "Calibri"
        font.size = Pt(8)

    def __getattr__(self, attr):
        return getattr(self.doc, attr)

    def _parse_patients(self):
        """Parse the patients from the Word document table into a PatientList."""
        pt_list = PatientList(home_ward=self.team.home_ward)

        for row in self._handover_table.rows:
            # skip past the column headers
            if row.cells[0].text.lower() == "bed":
                continue

            # skip past the ward name headers or empty rows
            if row.cells[0].text == row.cells[1].text:
                continue

            pt = Patient.from_table_row(row)
            pt_list.append(pt)

        return pt_list

    def _update_patients(self):
        updated_list = PatientList(home_ward=self.team.home_ward)
        latest_careflow_pts = get_careflow_patients(team=self.team)

        for careflow_patient in latest_careflow_pts:
            if careflow_patient not in self.patients:
                # patient must be new to the team
                careflow_patient.is_new = True
                # TODO: set 'reason_for_admission' here
                # careflow_patient.reason_for_admission = get_reason_for_admission(careflow_patient)
                updated_list.append(careflow_patient)
            else:
                # patient must be on the original list, therefore merge the
                # jobs, EDD etc into the careflow patient. By merging into
                # the careflow_pt, we also ensure that their location is fully
                # up to date incase they were moved since yesterday
                careflow_patient.merge(self.patients[careflow_patient.nhs_number])
                updated_list.append(careflow_patient)

        updated_list.sort()
        self.patients = updated_list

    @property
    def _handover_table(self):
        """Return the underlying Word document table"""
        return HandoverTable(self.tables[0])

    def _update_handover_table(self):
        table = self._handover_table

        # clear the existing table back to the column headers
        table.clear()

        current_ward = None
        for pt in self.patients:
            if pt.location.ward != current_ward:
                # we're at the first patient on a different ward, therefore add
                # a subheader row with the name of the new ward
                table.add_ward_header_row(pt.location.ward)
                current_ward = pt.location.ward

            table.add_patient_row(pt)

        # apply formatting to the newly updated table
        table.format()

        # put the patient count in the footer
        footer = self.sections[0].footer
        footer.footer_distance = Inches(1)
        footer.paragraphs[
            0
        ].text = f"{self.patient_count} patients ({self.new_patient_count} new)\n\n\n\n"

    def update(self):
        """Update the HandoverList patient table.

        Comprises of 2 phases:
            1) Update the in-memory 'self.patients' PatientList using CareFlow +/- TrakCare
            2) Update the underlying table in the Word document with
               the updated 'self.patients'
        """
        self._update_patients()
        self._update_handover_table()

    @property
    def patient_count(self):
        return len(self.patients)

    @property
    def new_patient_count(self):
        return sum(patient.is_new for patient in self.patients)
