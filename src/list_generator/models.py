import logging
from collections import defaultdict
from datetime import date, datetime, timedelta
from typing import Union

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn
from docx.shared import Cm, Pt
from docx.table import Table

from .. import shared_enums
from ..front_end.app import db
from ..shared_models import Patient
from ..utils import pluralise

logger = logging.getLogger()


class HandoverList:
    """The primary object representing the Word document containing the team's list of patients."""

    def __init__(self, team, file, filename):
        logger.debug("Instantiating HandoverList for %r using '%s' as a base list", team, filename)
        self.doc = Document(docx=file)
        self.team = team
        self.patients = self._parse_patients()
        logger.debug("Patients parsed from input handover list: %s", self.patients)
        if len(self.patients):
            logger.debug(
                "Parsed %d %s from the input handover list:\n\t%s",
                len(self.patients),
                pluralise("patient", len(self.patients)),
                "\n\t".join(repr(patient) for patient in self.patients),
            )
        else:
            logger.debug("No patients found on the input handover list")

        logger.debug("Applying default formatting to the Word document")
        self._apply_default_formatting()

    def __getattr__(self, attr):
        """Pass any unresolved attribute accesses down to the underlying docx.Document instance."""
        return getattr(self.doc, attr)

    def _parse_patients(self):
        """Parse the patients from the Word document table into a PatientList."""
        patient_list = PatientList(home_ward=self.team.home_ward)

        for row in self._handover_table.rows:
            # skip past the column headers
            if row.cells[0].text.lower() == "bed":
                continue

            # skip past the ward name headers or empty rows
            if row.cells[0].text == row.cells[1].text:
                continue
            try:
                pt = Patient.from_table_row(row)
            except ValueError as e:
                logger.exception(e)
                raise

            patient_list.append(pt)

        return patient_list

    @property
    def total_patient_count(self) -> int:
        return len(self.patients)

    @property
    def new_patient_count(self) -> int:
        return sum(patient.is_new for patient in self.patients)

    def get_trakcare_patients(self):
        """Fetch and return a list of patients from TrakCare under the current team."""
        team_name = self.team.name
        allowed_wards = [ward.value for ward in shared_enums.Ward]
        patients = (
            db.session.query(Patient)
            .filter(db.and_(Patient.team == team_name.value), Patient.ward.in_(allowed_wards))
            .all()
        )
        return patients

    @property
    def _handover_table(self):
        """Return the underlying Word document table."""
        return HandoverTable(self.tables[0])

    def _apply_default_formatting(self):
        style = self.styles["Normal"]
        font = style.font
        font.name = "Calibri"
        font.size = Pt(8)
        sections = self.doc.sections

        # set narrow margins
        for section in sections:
            section.top_margin = Cm(1.25)
            section.bottom_margin = Cm(1.25)
            section.left_margin = Cm(1.25)
            section.right_margin = Cm(1.25)

    def _update_patients(self) -> None:
        """Bring the internal PatientList object up to date with TrakCare."""
        updated_list = PatientList(home_ward=self.team.home_ward)
        logger.debug("Fetching patients from TrakCare")
        current_trakcare_patients = self.get_trakcare_patients()
        if len(current_trakcare_patients):
            logger.debug(
                "Found %d %s on TrakCare:\n\t%s",
                len(current_trakcare_patients),
                pluralise("patient", len(current_trakcare_patients)),
                "\n\t".join(repr(patient) for patient in current_trakcare_patients),
            )
        else:
            logger.debug("No patients found on TrakCare")

        # use the TrakCare patient list as the canonical source for which patients
        # are presently under the team. Note that any patients who were on the original
        # input handover list will be dropped at this point.
        for trakcare_patient in current_trakcare_patients:
            if trakcare_patient not in self.patients:
                # patient must be new to the team
                trakcare_patient.is_new = True
                updated_list.append(trakcare_patient)
            else:
                # patient must be on the original list, therefore merge the jobs,
                # EDD etc into the TrakCare patient. By merging into the TrakCare
                # patient, we also ensure that their location is fully up to date
                # incase they were moved since the previous list was produced.
                trakcare_patient.merge(self.patients[trakcare_patient.nhs_number])
                updated_list.append(trakcare_patient)

        logger.debug("Sorting the updated list")
        updated_list.sort()
        self.patients = updated_list
        if len(self.patients):
            logger.debug(
                "The updated list now has %d %s:\n\t%s",
                len(self.patients),
                pluralise("patient", len(self.patients)),
                "\n\t".join(repr(patient) for patient in self.patients),
            )
        else:
            logger.debug("No patients present on the updated list")

    def _update_handover_table(self) -> None:
        """Create a fresh Word table with the current PatientList."""
        self._handover_table.update(self.patients)

    def _update_list_metadata(self) -> None:
        """Update additonal metadata, such as the footer."""
        # put the patient count in the footer
        footer = self.sections[0].footer
        footer.footer_distance = Cm(1.25)
        metadata_text = (
            f"{self.total_patient_count} {pluralise('patient', self.total_patient_count)} "
            f"({self.new_patient_count} new)\t"
            f"\t"
            f"Generated at {datetime.now():%H:%M %d/%m/%Y}"
        )
        footer.paragraphs[0].text = metadata_text
        logger.debug("Added metadata to the handover list: %r", metadata_text)

    def update(self) -> None:
        """Update the HandoverList patient table.

        Comprises of 3 phases:
            1) Update the in-memory 'self.patients' PatientList using TrakCare
            2) Update the underlying table in the Word document with the updated 'self.patients'
            3) Update any addtional metadata pertaining to the HandoverList e.g. footer information
        """
        self._update_patients()
        self._update_handover_table()
        self._update_list_metadata()

    def __repr__(self):
        return (
            f"<{self.__class__.__name__}("
            f"team={self.team}, total_patient_count={self.total_patient_count}"
            f")>"
        )


class HandoverTable:
    def __init__(self, table: Table):
        self._table = table
        self._new_patient_indices = set()

    def __getattr__(self, attr):
        """Pass any unfound attributes down to the underlying Table object."""
        return getattr(self._table, attr)

    def update(self, patients):
        """Create a freshly updated table with the given patient list."""
        logger.debug("Generating a fresh table in the Word document")

        # clear the existing table back to the column headers
        self.clear()

        current_ward = None
        for patient in patients:
            if patient.location.ward != current_ward:
                # we're at the first patient on a different ward, therefore add
                # a subheader row with the name of the new ward
                current_ward = patient.location.ward
                self.add_ward_header_row(current_ward)

            self.add_patient_row(patient)
            logger.debug("%r added to the table", patient)

        if not patients:
            logger.debug("No patients added to the table")

        # apply formatting to the newly updated table
        logger.debug("Applying table formatting to the Word document")
        self.format()

    def clear(self, keep_headers: bool = True) -> None:
        logger.debug("Removing old rows from original patient list table")
        for row in self.rows[keep_headers:]:
            tr = row._tr
            self._tbl.remove(tr)

    def add_full_width_row(self):
        new_row = self.add_row()
        new_row.cells[0].merge(new_row.cells[-1])
        return new_row

    def add_ward_header_row(self, ward: shared_enums.Ward) -> None:
        ward_header_row = self.add_full_width_row()
        ward_header_row.cells[0].text = ward.value

    def add_patient_row(self, patient: Patient) -> None:
        if patient.is_birthday:
            birthday_row = self.add_full_width_row()
            birthday_row.cells[0].text = f"Happy birthday {patient.forename}! {patient.age} today!"

        new_patient_row = self.add_row()
        new_patient_row.patient = patient

        new_patient_row.cells[0].text = patient.bed
        new_patient_row.cells[1].text = patient.patient_details
        new_patient_row.cells[2].text = patient.reason_for_admission or ""
        if not patient.is_new:
            # new patients won't have any of these attributes on them
            new_patient_row.cells[3].text = patient.jobs
            new_patient_row.cells[4].text = patient.edd
            new_patient_row.cells[5].text = patient.ds
            new_patient_row.cells[6].text = patient.tta
            new_patient_row.cells[7].text = patient.bloods
        else:
            self._new_patient_indices.add(new_patient_row._index)

    def format(self) -> None:
        # center table within the page
        self.alignment = WD_TABLE_ALIGNMENT.CENTER

        # set column width for the patient details
        # 3cm is a good default width to fit dob and age on one line
        for cell in self.columns[1].cells:
            cell.width = Cm(3)

        # format the rows and cells
        for row in self.rows:
            # format the column headers and set header row to repeat
            if row.cells[0].text.lower() == "bed":
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.underline = True
                tr = row._tr
                trPr = tr.get_or_add_trPr()
                tblHeader = OxmlElement("w:tblHeader")
                tblHeader.set(qn("w:val"), "true")
                trPr.append(tblHeader)

            # format the ward headers
            if row.cells[0].text == row.cells[1].text:
                shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="EEEEEE"/>')
                for cell in row.cells:
                    cell._tc.get_or_add_tcPr().append(shading_elm)
                    # keep this header row on the same page as the next row
                    for paragraph in cell.paragraphs:
                        paragraph.paragraph_format.keep_with_next = True
                if "birthday" in row.cells[0].text:
                    for cell in row.cells:
                        for para in cell.paragraphs:
                            for run in para.runs:
                                run.italic = True

            # format new patients as bold
            if row._index in self._new_patient_indices:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        for run in para.runs:
                            run.bold = True

            # format all cells in the table
            for cell in row.cells:
                cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                for paragraph in cell.paragraphs:
                    formatter = paragraph.paragraph_format
                    formatter.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                    formatter.space_before = Pt(2)
                    formatter.space_after = Pt(2)

            # prevent table rows from splitting across page breaks
            no_split = parse_xml(f"<w:cantSplit {nsdecls('w')}/>")
            row._tr.get_or_add_trPr().append(no_split)


class PatientList:
    def __init__(self, home_ward: shared_enums.Ward):
        self.home_ward: shared_enums.Ward = home_ward
        # a mapping of {nhs_number: Patient}
        self._patient_mapping: dict = {}

    def append(self, patient: "Patient") -> None:
        self[patient] = patient

    def extend(self, patients: list) -> None:
        for patient in patients:
            self.append(patient)

    def sort(self) -> None:
        """Sort the patient list in place.

        Applies the following sorting strategy:
            1) Sort the home ward patients to the top of the list
            2) Sort the remaining outlier wards alphabetically
            3) Sort the patients within each ward by bed
        """
        sorted_dict = {}
        grouped_dict = defaultdict(list)
        # create a mapping of {ward: [Patient]}
        for patient in self:
            grouped_dict[patient.location.ward].append(patient)

        # sort the patients by bed on a per-ward basis
        for ward, pts in grouped_dict.items():
            grouped_dict[ward] = sorted(pts, key=lambda patient: patient.location._bed_sort)

        # populate the new list with the Home Ward patients first
        for patient in grouped_dict.pop(self.home_ward, []):
            sorted_dict[patient.nhs_number] = patient

        # then populate the new list with the outlier patients sorted by ward alphabetically
        for ward in sorted(grouped_dict.keys(), key=lambda k: k.value):
            for patient in grouped_dict[ward]:
                sorted_dict[patient.nhs_number] = patient

        self._patient_mapping = sorted_dict

    @property
    def patients(self):
        return list(self)

    @patients.setter
    def patients(self, value):
        raise AttributeError(f"{self.__class__.__name__}.patients is a read-only property.")

    def __getitem__(self, key: Union[str, "Patient"]) -> "Patient":
        """Return a Patient from the list with a given Patient or NHS number."""
        if isinstance(key, Patient):
            key = key.nhs_number
        try:
            return self._patient_mapping[key]
        except KeyError:
            raise KeyError(f"No patient with the NHS number '{key}' found in the list")

    def __setitem__(self, _: str, value: "Patient"):
        if not isinstance(value, Patient):
            raise TypeError

        self._patient_mapping[value.nhs_number] = value

    def __contains__(self, key: Union[str, "Patient"]) -> bool:
        """Assert whether a given patient (where type(key) == Patient) is in the list."""
        if isinstance(key, Patient):
            key = key.nhs_number
        return key in self._patient_mapping

    def __iter__(self):
        return iter(self._patient_mapping.values())

    def __len__(self):
        return len(self._patient_mapping)

    def __repr__(self):
        return (
            f"<{self.__class__.__name__}("
            f"home_ward={self.home_ward.value}, "
            f"total_patient_count={len(self)}, "
            f"new_patient_count={len([pt for pt in self if pt.is_new])})>"
        )
